import docx
import datetime

i = datetime.datetime.now()

doc=docx.Document('E:\mp\通源小学网络安全事件信息报告表.docx')
newcont = '报告时间：%s年%s月%s日%s:%s分         签发人：杨海中' % (i.year, i.month, i.day-1, i.hour, i.minute)
run = doc.paragraphs[2].add_run(newcont)
run.font.name = 'simhei'
run.font.size = 160000
doc.save('%s月%s日通源小学网络安全事件信息报告表.docx' % (i.month, i.day-1))
