import smtplib
from email.mime.multipart import MIMEMultipart    
#from email.mime.text import MIMEText    
#from email.mime.image import MIMEImage 
from email.header import Header
import datetime
import docx

i = datetime.datetime.now()

doc=docx.Document('通源小学网络安全事件信息报告表.docx')
doc.paragraphs[2].text = '报告时间：%s年%s月%s日%s:%s分         签发人：杨海中' % (i.year, i.month, i.day, i.hour, i.minute)
run = doc.paragraphs[2]
#run.font.name = '微软雅黑'
#run.font.size = 140000
doc.save('%s月%s日通源小学网络安全事件信息报告表.docx' % (i.month, i.day-1))

#设置smtplib所需的参数
smtpserver = 'smtp.qq.com,465'
username = 'yhzhm@qq.com'
password = 'denpxueeegtfbjaf'
sender = 'yhzhm@qq.com'
#receiver = ['1022492@qq.com']
receiver = ['49186016@qq.com']

#subject = '通源小学无网络安全问题'
subject = '%s月%s日 通源小学无网络安全问题' % (i.month, i.day-1 )
subject = Header(subject,'utf-8').encode()

msg = MIMEMultipart('mixed') 
msg['Subject'] = subject
msg['From'] = sender
msg['To'] = receiver[0]
#msg['Date']='2018-6-3'


try:
    smtp = smtplib.SMTP_SSL() 
    smtp.connect('smtp.qq.com',465)    #465或587
    smtp.login(username, password) 
    smtp.sendmail(sender, receiver, msg.as_string()) 
    smtp.quit()
    print("Sucess!")
except smtplib.SMTPException as e:
    print("Falied,%s" %e)
